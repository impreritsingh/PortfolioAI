import os
import tempfile
from typing import Dict, Any
import io
import PyPDF2
from groq import Groq
import docx
import gradio as gr
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Initialize FastAPI app
app = FastAPI(title="Resume Scoring API")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Groq client
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
if not GROQ_API_KEY:
    print("Warning: GROQ_API_KEY not found in environment variables. Create a .env file with GROQ_API_KEY=your_api_key")

groq_client = Groq(api_key=GROQ_API_KEY)
LLM_MODEL = "meta-llama/llama-4-maverick-17b-128e-instruct"  # You can change this to any model Groq supports

class ResumeScore(BaseModel):
    score: int
    feedback: Dict[str, Any]

async def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file bytes"""
    try:
        with PyPDF2.PdfReader(io.BytesIO(file_content)) as pdf_reader:
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from PDF: {str(e)}")

async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file bytes"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from DOCX: {str(e)}")

async def extract_text_from_resume(file: UploadFile) -> str:
    """Extract text from resume file based on its extension"""
    file_content = await file.read()
    file_extension = file.filename.split(".")[-1].lower()
    
    if file_extension == "pdf":
        return await extract_text_from_pdf(file_content)
    elif file_extension in ["docx", "doc"]:
        return await extract_text_from_docx(file_content)
    elif file_extension == "txt":
        return file_content.decode("utf-8")
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF, DOCX, or TXT file.")

async def score_resume(resume_text: str) -> ResumeScore:
    """Use Groq's LLM to score and provide feedback on the resume"""
    if not GROQ_API_KEY:
        # Provide mock response for testing without API key
        return ResumeScore(
            score=75,
            feedback={
                "strengths": ["Strong education section", "Good experience details"],
                "weaknesses": ["Missing quantifiable achievements", "Too verbose"],
                "improvements": ["Add metrics to achievements", "Be more concise", "Consider adding skills section"]
            }
        )
    
    try:
        prompt = f"""
        I have a resume that I'd like you to evaluate. Please analyze it carefully and provide:
        
        1. A numerical score from 0-100 reflecting the overall quality
        2. Key strengths (at least 3)
        3. Areas for improvement (at least 3)
        4. Specific, actionable suggestions for making this resume more effective
        
        Focus on factors like: clarity, quantifiable achievements, relevance, formatting consistency, use of action verbs, 
        and overall impact. Please format your response as a JSON object with the following structure:
        {{
            "score": (numerical score),
            "feedback": {{
                "strengths": [(list of strengths)],
                "weaknesses": [(list of weaknesses)],
                "improvements": [(list of specific suggestions)]
            }}
        }}
        
        Here's the resume text:
        {resume_text}
        """
        
        response = groq_client.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are an expert resume reviewer with years of HR and recruiting experience."},
                {"role": "user", "content": prompt}
            ],
            model=LLM_MODEL,
            response_format={"type": "json_object"}
        )
        
        result = response.choices[0].message.content
        import json
        result_json = json.loads(result)
        
        return ResumeScore(
            score=result_json["score"],
            feedback=result_json["feedback"]
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error scoring resume: {str(e)}")

@app.post("/score-resume/", response_model=ResumeScore)
async def score_resume_api(file: UploadFile = File(...)):
    """API endpoint to score a resume"""
    resume_text = await extract_text_from_resume(file)
    return await score_resume(resume_text)

# Create Gradio interface
def gradio_score_resume(file):
    """Function for Gradio interface to score a resume"""
    try:
        # Get file content from Gradio file object
        file_content = file.name
        file_extension = file.name.split(".")[-1].lower()
        
        if file_extension == "pdf":
            with open(file_content, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
        elif file_extension in ["docx", "doc"]:
            doc = docx.Document(file_content)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_extension == "txt":
            with open(file_content, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            return None, "Unsupported file format. Please upload PDF, DOCX, or TXT file."
        
        # Pass directly to async function and run in event loop
        import asyncio
        result = asyncio.run(score_resume(text))
        
        # Format the response for display
        strengths = "\n".join([f"- {s}" for s in result.feedback["strengths"]])
        weaknesses = "\n".join([f"- {s}" for s in result.feedback["weaknesses"]])
        improvements = "\n".join([f"- {s}" for s in result.feedback["improvements"]])
        
        feedback_text = f"""
        ## Resume Score: {result.score}/100
        
        ### Strengths:
        {strengths}
        
        ### Areas for Improvement:
        {weaknesses}
        
        ### Specific Recommendations:
        {improvements}
        """
        
        return result.score, feedback_text
    
    except Exception as e:
        return None, f"Error processing resume: {str(e)}"

# Create Gradio interface
demo = gr.Interface(
    fn=gradio_score_resume,
    inputs=gr.File(label="Upload your resume (PDF, DOCX, or TXT)"),
    outputs=[
        gr.Number(label="Resume Score (0-100)"),
        gr.Markdown(label="Detailed Feedback")
    ],
    title="Resume Scorer",
    description="Upload your resume to get a score and personalized improvement suggestions.",
    examples=[],
    allow_flagging="never",
    theme=gr.themes.Soft()
)

# Mount the Gradio app to FastAPI
app = gr.mount_gradio_app(app, demo, path="/")

# For Hugging Face Spaces compatibility
if os.environ.get("SPACE_ID"):
    print("Running on Hugging Face Spaces")
    # In Hugging Face Spaces, we need to use a specific port
    port = int(os.environ.get("PORT", 7860))
    
    if __name__ == "__main__":
        import uvicorn
        uvicorn.run(app, host="0.0.0.0", port=port)
else:
    # Run normally for local development
    if __name__ == "__main__":
        import uvicorn
        uvicorn.run(app, host="0.0.0.0", port=8000)